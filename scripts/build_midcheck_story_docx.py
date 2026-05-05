#!/usr/bin/env python3
"""Build the professor-facing midcheck DOCX report."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "midcheck" / "ETRI_중간점검_변수스토리_20260505.docx"

ACCENT = "1F4E5F"
ACCENT_LIGHT = "EAF3F6"
GRID = "D9E2E7"
MUTED = RGBColor(95, 103, 111)
TEXT = RGBColor(32, 36, 40)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = GRID, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_table(table, header=True) -> None:
    table.style = "Table Grid"
    table.autofit = False
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = TEXT
        if header and r_idx == 0:
            for cell in row.cells:
                set_cell_shading(cell, ACCENT_LIGHT)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(20, 72, 88)


def set_width(cell, width: float) -> None:
    cell.width = Inches(width)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_run(paragraph, text: str, *, bold=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def add_note(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_width(cell, 6.5)
    set_cell_shading(cell, "F6FAFB")
    set_cell_border(cell, color="BFD4DB", size="8")
    set_cell_margins(cell, top=140, bottom=140, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_run(p, title, bold=True, color=RGBColor(20, 72, 88), size=11)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    add_run(p2, body, color=TEXT, size=10)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_after = Pt(5)
        add_run(p, item, size=10.5)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    title = styles["Title"]
    title.font.name = "Arial"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    title.font.size = Pt(21)
    title.font.bold = True
    title.font.color.rgb = RGBColor(18, 69, 84)
    title.paragraph_format.space_after = Pt(4)

    for name, size in [("Heading 1", 15), ("Heading 2", 12.5)]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(18, 69, 84)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(0)
        add_run(hp, "ETRI 인간이해 대회 중간점검", color=MUTED, size=8.5)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(fp, "2026-05-05", color=MUTED, size=8.5)


def add_score_table(doc: Document) -> None:
    rows = [
        ["시점", "파일/전략", "Public", "해석"],
        ["4/25", "LGB target-wise histmix", "0.5960566585", "일 단위 라이프로그와 타깃별 LightGBM 기준선"],
        ["4/29", "stable tuned", "0.5956332255", "타깃별 안정 피처와 파라미터 조정"],
        ["5/01", "temporal prior", "0.5886545849", "같은 사람의 가까운 날짜 정보가 강하게 작동"],
        ["5/02", "temporal targetwise", "0.5863944910", "타깃별 시간 정보 반영"],
        ["5/04", "lgb_temporal_s4b650", "0.5829008297", "현재 best. S4의 시간 연속성 조정이 유효"],
        ["5/05", "guard 정책 전환", "-", "추가 후처리는 엄격한 조건 통과 시만 제출"],
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    widths = [0.65, 1.65, 1.1, 3.1]
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            cell = table.cell(r, c)
            set_width(cell, widths[c])
            cell.text = text
    style_table(table)


def add_target_table(doc: Document) -> None:
    rows = [
        ["타깃", "의미", "내가 연결해서 본 라이프로그 신호"],
        ["Q1", "기상 직후 주관적 수면 질", "수면 중 심박 안정성, 새벽 화면 사용, 개인 기준 대비 변화"],
        ["Q2", "취침 전 피로도", "낮 동안 걸음/활동량, 이동 범위, 심박 부담, 충전/휴대폰 사용 루틴"],
        ["Q3", "취침 전 스트레스", "활동 전환, 화면 세션 밀도, 이동·사회적 노출, 개인 기준 대비 이탈"],
        ["S1", "총 수면 시간 권장 기준 충족", "수면 시간대 활동/화면/조도와 안정적인 야간 패턴"],
        ["S2", "수면 효율 권장 기준 충족", "새벽 활동, 화면 켜짐, 심박 변동, 야간 움직임"],
        ["S3", "입면 지연 권장 기준 충족", "잠들기 전 화면·활동·조도, 저녁/새벽 행동 전환"],
        ["S4", "수면 중 각성 시간(WASO) 권장 기준 충족", "밤중 화면·걸음·활동·심박 변화, 같은 사람의 인접 날짜 연속성"],
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    widths = [0.55, 2.15, 3.8]
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            cell = table.cell(r, c)
            set_width(cell, widths[c])
            cell.text = text
    style_table(table)


def add_feature_story_table(doc: Document) -> None:
    rows = [
        ["스토리 축", "사용한 변수/조합", "왜 봤는가"],
        [
            "개인 기준",
            "subject z-score, target lag1, 최근 rolling target encoding",
            "Q1~Q3는 개인 평균보다 좋은지/나쁜지로 만들어진 라벨이라 절대값보다 '평소의 나와 다른가'가 중요하다.",
        ],
        [
            "수면 시간대",
            "00~09시 심박 평균·분위수·RMSSD·spike, 수면 중 걸음·활동·화면·조도",
            "실제 수면 건강 지표 S1~S4는 밤 사이의 안정성, 각성, 움직임에서 직접적인 단서가 나온다고 판단했다.",
        ],
        [
            "밤중 방해",
            "overnight/evening screen share + night steps, light zero/share, screen session count",
            "새벽에 화면이 켜지거나 움직임이 많으면 수면 연속성이 깨졌을 가능성이 있어 S2/S3/S4와 Q1에 연결했다.",
        ],
        [
            "주간 피로 누적",
            "step sum, active min, GPS extent, WiFi/BLE unique count, HR mean/rest q10",
            "낮의 활동 부하와 이동량은 취침 전 피로(Q2), 스트레스(Q3), 다음 날 수면 질(Q1)에 영향을 준다고 보았다.",
        ],
        [
            "생활 루틴",
            "screen session density, charge min/share, app/speech share, modal coverage",
            "휴대폰 사용과 충전 패턴은 하루의 리듬, 잠들기 전 준비 상태, 데이터 관측 신뢰도를 함께 설명한다.",
        ],
        [
            "시간 연속성",
            "same-subject temporal prior, S4 beta ladder",
            "수면 문제는 하루 단위로 완전히 독립적이지 않으므로 가까운 날짜의 같은 사람 패턴을 조심스럽게 반영했다.",
        ],
    ]
    table = doc.add_table(rows=len(rows), cols=3)
    widths = [1.15, 2.35, 3.0]
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            cell = table.cell(r, c)
            set_width(cell, widths[c])
            cell.text = text
    style_table(table)


def main() -> None:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(title, "ETRI 인간이해 대회 중간점검", bold=True, size=21, color=RGBColor(18, 69, 84))

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    add_run(
        subtitle,
        "점수 개선 기록보다 중요한 것: 라이프로그 변수에 의미를 부여한 설계 스토리",
        color=MUTED,
        size=11,
    )

    meta = doc.add_table(rows=3, cols=2)
    meta_rows = [
        ("작성일", "2026-05-05"),
        ("현재 제출 기준", "lgb_temporal_s4b650.csv / Public 0.5829008297"),
        ("데이터 규모", "train 450행, test 250행, subject 10명, 타깃 7개(Q1~Q3, S1~S4)"),
    ]
    for r, (k, v) in enumerate(meta_rows):
        meta.cell(r, 0).text = k
        meta.cell(r, 1).text = v
        set_width(meta.cell(r, 0), 1.3)
        set_width(meta.cell(r, 1), 5.2)
    style_table(meta, header=False)
    for row in meta.rows:
        set_cell_shading(row.cells[0], ACCENT_LIGHT)
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(20, 72, 88)

    add_note(
        doc,
        "교수님 코멘트 반영 방향",
        "단순히 rolling 변수를 많이 만든 것이 아니라, '왜 이 변수를 봤는가'와 '이 변수와 저 변수를 결합하면 어떤 생활 상태를 뜻하는가'가 보이도록 정리했다.",
    )

    doc.add_heading("1. 문제를 이렇게 해석했다", level=1)
    p = doc.add_paragraph()
    add_run(
        p,
        "이번 대회는 스마트폰, 웨어러블, 위치, 조도, 활동 로그를 이용해 다음 날의 주관적 수면 상태와 수면 센서 기반 권장 기준 충족 여부를 예측하는 문제다. 핵심은 개인마다 절대적인 생활 패턴이 다르다는 점이다. 그래서 같은 걸음 수, 같은 화면 사용 시간이라도 어떤 사람에게는 평소 수준이고, 다른 사람에게는 비정상적인 신호일 수 있다.",
        size=10.5,
    )
    add_bullets(
        doc,
        [
            "Q1~Q3는 개인 평균 대비 좋은 상태인지 나쁜 상태인지로 해석했다.",
            "S1~S4는 실제 수면 센서 기준의 수면 시간, 효율, 입면 지연, 수면 중 각성 문제로 해석했다.",
            "따라서 '전체 평균에서 높은가'보다 '그 사람의 평소와 얼마나 다른가'를 중심에 두었다.",
        ],
    )

    doc.add_heading("2. 타깃별로 본 생활 신호", level=1)
    add_target_table(doc)

    doc.add_heading("3. 변수 설계 스토리", level=1)
    add_feature_story_table(doc)

    doc.add_heading("4. 모델링은 의미를 보존하는 방향으로 했다", level=1)
    add_bullets(
        doc,
        [
            "기본 모델은 LightGBM이다. 데이터가 450행으로 작고 피처가 많아, 비선형 관계를 잡되 과적합을 통제하기 쉬운 모델이 필요했다.",
            "모든 타깃을 한 모델처럼 다루지 않고 target-wise로 학습했다. S4와 Q1은 어려운 타깃이라 규제를 강하게 주고, S1은 비교적 쉬운 타깃이라 모델 용량을 조금 더 허용했다.",
            "단일 점수만 보지 않고 OOF, public 점수, test drift를 함께 보며 제출 여부를 판단했다.",
            "최근에는 새 CSV를 무조건 만드는 대신, 의미 있는 개선 근거가 있는 후보만 제출하는 guard 정책으로 전환했다.",
        ],
    )

    doc.add_page_break()
    doc.add_heading("5. 실험 흐름과 현재 위치", level=1)
    add_score_table(doc)
    p = doc.add_paragraph()
    add_run(
        p,
        "현재 best는 S4의 시간 연속성을 반영한 lgb_temporal_s4b650.csv이다. 다만 S4만 계속 밀어붙이는 방식은 이후 public에서 더 좋아지지 않았기 때문에, 다음 단계는 S4 단독 조정보다 non-S4 타깃까지 같이 설명할 수 있는 feature-stable 재학습으로 잡았다.",
        size=10.5,
    )

    doc.add_heading("6. 다음 중간점검까지의 계획", level=1)
    add_bullets(
        doc,
        [
            "새 제출 파일은 guard를 통과할 때만 만든다. 통과하지 못하면 제출하지 않는 것이 전략이다.",
            "1순위는 안정적인 피처 subset을 타깃별로 다시 고르는 것이다. 많은 변수를 쓰는 것보다 반복 seed와 fold에서 계속 살아남는 변수를 우선한다.",
            "2순위는 시간 prior를 다시 설계하되 S4만 단독으로 움직이지 않는 것이다. Q1~Q3와 S1~S3까지 함께 설명되는 변화만 후보로 본다.",
            "XGBoost나 다른 모델은 다양성 후보로만 사용한다. raw OOF와 guarded blend가 모두 통과하지 않으면 제출 파일로 만들지 않는다.",
        ],
    )

    doc.add_heading("한 줄 결론", level=1)
    add_note(
        doc,
        "중간점검 결론",
        "이번 작업의 핵심은 '수면은 개인의 평소 리듬에서 벗어나는 순간 흔들린다'는 가설이다. 그래서 개인 기준, 수면 시간대 신호, 밤중 방해, 주간 피로 누적, 시간 연속성을 중심으로 변수를 만들고 모델을 설계했다.",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
